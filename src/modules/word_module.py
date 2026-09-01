import os
import re
import zipfile
from src.core.base_module import BaseDocumentModule
from src.core.registry import ModuleRegistry

class WordModule(BaseDocumentModule):
    @property
    def name(self) -> str:
        return "Word"

    @property
    def file_extensions(self) -> list[str]:
        return [".docx"]

    @property
    def required_dependencies(self) -> list[str]:
        return ["python-docx", "Pillow"]

    def load_to_markdown(self, file_path: str) -> str:
        """Extracts Word .docx to clean Markdown text, preserving tables, headings, bold/italic styles, and lists."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            import docx
            from docx.oxml.text.paragraph import CT_P
            from docx.oxml.table import CT_Tbl
            from docx.text.paragraph import Paragraph
            from docx.table import Table

            if not zipfile.is_zipfile(file_path):
                return "⚠️ File Validation Error: Invalid file. Please select a valid DOCX document."

            doc = docx.Document(file_path)
            parts = []

            def iter_block_items(parent):
                parent_elm = parent.element.body
                for child in parent_elm.iterchildren():
                    if isinstance(child, CT_P):
                        yield Paragraph(child, parent)
                    elif isinstance(child, CT_Tbl):
                        yield Table(child, parent)

            from docx.oxml.ns import qn
            from src.core.converters import wrap_text_style
            from src.services.media_asset_manager import MediaAssetManager
            asset_mgr = MediaAssetManager()

            def extract_and_get_link(run) -> str:
                xml = run._element.xml
                rId_match = re.search(r'(?:r:embed|embed)="([^"]+)"', xml)
                if rId_match:
                    rId = rId_match.group(1)
                    try:
                        image_part = doc.part.related_parts[rId]
                        if hasattr(image_part, "image"):
                            image_bytes = image_part.image.blob
                            ext = "png"
                            if hasattr(image_part.image, "ext") and image_part.image.ext:
                                ext = image_part.image.ext
                            elif hasattr(image_part.image, "content_type") and image_part.image.content_type:
                                ct = image_part.image.content_type
                                if "/" in ct:
                                    ext = ct.split("/")[1]
                            
                            filename = f"image_{rId}.{ext}"
                            virtual_uri = asset_mgr.register_image(image_bytes, filename)

                            # Extract extent width if available in Word Drawing XML
                            width_str = ""
                            cx_match = re.search(r'(?:wp:extent|a:ext)\s+[^>]*?cx=["\'](\d+)["\']', xml)
                            if cx_match:
                                try:
                                    cx_emus = int(cx_match.group(1))
                                    # Standard printable width on A4 is 6.0 inches (5,486,400 EMUs)
                                    width_inches = cx_emus / 914400.0
                                    width_pct = int(round((width_inches / 6.0) * 100))
                                    if 10 <= width_pct <= 92:
                                        # Snap to nearest 5% or 25% preset if close
                                        if abs(width_pct - 25) <= 3:
                                            width_pct = 25
                                        elif abs(width_pct - 50) <= 3:
                                            width_pct = 50
                                        elif abs(width_pct - 75) <= 3:
                                            width_pct = 75
                                        width_str = f'{width_pct}%'
                                except Exception:
                                    pass

                            # Determine paragraph alignment if present
                            para_align = ""
                            if block.alignment is not None:
                                from docx.enum.text import WD_ALIGN_PARAGRAPH
                                if block.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                                    para_align = "center"
                                elif block.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                                    para_align = "right"

                            if para_align or width_str:
                                w_attr = f' width="{width_str}"' if width_str else ""
                                img_tag = f'<img src="{virtual_uri}" alt="image"{w_attr} />'
                                if para_align:
                                    return f'<p align="{para_align}">{img_tag}</p>'
                                return img_tag

                            return f"![image]({virtual_uri})"
                    except Exception as e:
                        print(f"[DEBUG] Failed to extract run image with rId {rId}: {e}")
                return "[image]"

            for block in iter_block_items(doc):
                if isinstance(block, Paragraph):
                    style_name = (block.style.name or "").lower() if block.style else ""
                    
                    def run_contains_image(run):
                        xml = getattr(run._element, "xml", "")
                        return "<w:drawing" in xml or "<w:pict" in xml or "<a:blip" in xml

                    para_parts = []
                    for child in block._element:
                        if child.tag == qn('w:r'):
                            run = docx.text.run.Run(child, block)
                            if run_contains_image(run):
                                para_parts.append(extract_and_get_link(run))
                                continue
                            if not run.text:
                                continue
                            formatted = wrap_text_style(
                                run.text,
                                bold=run.bold,
                                italic=run.italic,
                                strike=run.font.strike,
                                underline=run.font.underline
                            )
                            para_parts.append(formatted)
                        elif child.tag == qn('w:hyperlink'):
                            r_id = child.get(qn('r:id'))
                            url = ""
                            if r_id:
                                try:
                                    url = block.part.rels[r_id].target_ref
                                except Exception:
                                    pass
                            link_parts = []
                            for sub_child in child:
                                  if sub_child.tag == qn('w:r'):
                                     run = docx.text.run.Run(sub_child, block)
                                     if run_contains_image(run):
                                         link_parts.append(extract_and_get_link(run))
                                         continue
                                     if not run.text:
                                         continue
                                     formatted = wrap_text_style(
                                         run.text,
                                         bold=run.bold,
                                         italic=run.italic,
                                         strike=run.font.strike,
                                         underline=run.font.underline
                                     )
                                     link_parts.append(formatted)
                            link_text = "".join(link_parts)
                            if link_text:
                                if url:
                                    para_parts.append(f"[{link_text}]({url})")
                                else:
                                    para_parts.append(link_text)
                    
                    para_text = "".join(para_parts).strip()
                    if not para_text:
                        if block.text.strip():
                            para_text = block.text.strip()
                        else:
                            parts.append("")
                            continue
                    
                    is_heading = style_name.startswith("heading") or re.match(r"^(đề mục|tiêu đề)\s*\d", style_name)

                    if is_heading:
                        m = re.search(r"\d+", style_name)
                        level = int(m.group(0)) if m else 1
                        parts.append("#" * level + " " + para_text)
                    elif "bullet" in style_name:
                        m = re.search(r"\d+", style_name)
                        level = int(m.group(0)) if m else 1
                        indent = "  " * (level - 1)
                        parts.append(indent + "- " + para_text)
                    elif "number" in style_name or style_name.startswith("list"):
                        m = re.search(r"\d+", style_name)
                        level = int(m.group(0)) if m else 1
                        indent = "  " * (level - 1)
                        parts.append(indent + "1. " + para_text)
                    else:
                        parts.append(para_text)

                elif isinstance(block, Table):
                    table_parts = []
                    rows_data = []
                    for row in block.rows:
                        row_cells = []
                        for cell in row.cells:
                            cell_text = cell.text.strip().replace("\n", " ").replace("|", "\\|")
                            row_cells.append(cell_text)
                        rows_data.append(row_cells)
                    
                    if rows_data:
                        header = "| " + " | ".join(rows_data[0]) + " |"
                        sep = "| " + " | ".join("---" for _ in rows_data[0]) + " |"
                        table_parts.append(header)
                        table_parts.append(sep)
                        for row in rows_data[1:]:
                            table_parts.append("| " + " | ".join(row) + " |")
                        parts.append("\n".join(table_parts))

            return "\n\n".join(parts)
        except Exception as e:
            import sys
            print(f"[DEBUG] Custom Word parsing failed: {e}. Falling back to markitdown.", file=sys.stderr)
            try:
                from markitdown import MarkItDown
                md = MarkItDown()
                result = md.convert(file_path)
                if not result or not result.text_content:
                    return "*(Empty Word Document)*"
                return result.text_content
            except Exception as inner_e:
                raise RuntimeError(f"Word Ingestion Error: Failed to extract text layer from DOCX file. Detail: {str(inner_e)}")

    def save_from_markdown(self, markdown_content: str, out_path: str) -> str:
        """Converts Markdown text to formatted Word document."""
        import docx
        import re
        from src.core.converters import prepare_markdown_for_export
        markdown_content = prepare_markdown_for_export(markdown_content)

        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        FONT = "Arial"
        HEADING_SIZES = {1: 20, 2: 16, 3: 13, 4: 12, 5: 11, 6: 11}
        HEADING_COLORS = {i: "404040" for i in range(1, 7)}

        def set_font(run, size=11, bold=False, color=None):
            run.font.name = FONT
            run.font.size = Pt(size)
            run.font.bold = bold
            if color:
                run.font.color.rgb = RGBColor.from_string(color)

        def add_hyperlink_run(paragraph, text, url, size=11, bold=False, italic=False, strike=False, underline=True, color="0000FF"):
            part = paragraph.part
            r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)

            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')

            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), FONT)
            rFonts.set(qn('w:hAnsi'), FONT)
            rPr.append(rFonts)

            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(size * 2)))
            rPr.append(sz)

            if bold:
                b = OxmlElement('w:b')
                rPr.append(b)
            if italic:
                i = OxmlElement('w:i')
                rPr.append(i)
            if strike:
                strike_el = OxmlElement('w:strike')
                rPr.append(strike_el)
            if underline:
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'single')
                rPr.append(u)

            if color:
                c = OxmlElement('w:color')
                c.set(qn('w:val'), color)
                rPr.append(c)

            new_run.append(rPr)
            
            text_node = OxmlElement('w:t')
            text_node.text = text
            new_run.append(text_node)

            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)
            return hyperlink

        from src.core.converters import parse_inline

        def add_formatted_runs(paragraph, text, size=11, default_bold=False, default_color=None):
            segments = parse_inline(text, bold=default_bold)
            for seg in segments:
                if seg.url:
                    add_hyperlink_run(
                        paragraph, 
                        seg.text, 
                        seg.url, 
                        size=size, 
                        bold=seg.bold, 
                        italic=seg.italic, 
                        strike=seg.strike, 
                        underline=True, 
                        color="0000FF"
                    )
                else:
                    run = paragraph.add_run(seg.text)
                    run.font.name = "Consolas" if seg.code else FONT
                    run.font.size = Pt(size - 1) if seg.code else Pt(size)
                    run.font.bold = seg.bold
                    run.font.italic = seg.italic
                    run.font.strike = seg.strike
                    run.font.underline = seg.underline
                    if seg.code:
                        run.font.color.rgb = RGBColor(0xA5, 0x2A, 0x2A)
                    elif default_color:
                        run.font.color.rgb = RGBColor.from_string(default_color)

        def add_paragraph_with_font(doc, text, size=11, bold=False, color=None, style=None):
            p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
            add_formatted_runs(p, text, size=size, default_bold=bold, default_color=color)
            return p

        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from src.services.media_asset_manager import MediaAssetManager
        asset_mgr = MediaAssetManager()

        # Preprocess multi-line image wrapper blocks into single-line tokens so line loop processes them atomically
        markdown_content = re.sub(
            r'<(p|div|center)(?:\s+[^>]*?)?>\s*(<img\s+[^>]*?>|!\[[^\]]*\]\([^\)]+\))\s*</\1>',
            lambda m: m.group(0).replace("\r", " ").replace("\n", " "),
            markdown_content,
            flags=re.IGNORECASE
        )

        lines = markdown_content.splitlines()
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = FONT # type: ignore
        style.font.size = Pt(11) # type: ignore

        i = 0
        while i < len(lines):
            line = lines[i].rstrip("\r\n")

            # Image check (Fast string pre-check before running token parser)
            if "!" in line or "<img" in line or "<p" in line or "<div" in line or "<center" in line:
                from src.ui_flet.helpers.image_token_helper import find_all_image_tokens
                img_tokens = find_all_image_tokens(line)
                if img_tokens:
                    tok = img_tokens[0]
                    text_before = line[:tok.start].strip()
                    text_after = line[tok.end:].strip()

                    if text_before or text_after:
                        # Split inline image from surrounding text so no heading/paragraph text is lost
                        lines[i] = text_before if text_before else tok.raw_token
                        insert_offset = 1
                        if text_before:
                            lines.insert(i + insert_offset, tok.raw_token)
                            insert_offset += 1
                        if text_after:
                            lines.insert(i + insert_offset, text_after)
                        continue

                    src_url = tok.src
                    if src_url:
                        out_dir = os.path.dirname(out_path) if out_path else None
                        img_path = asset_mgr.resolve_uri(src_url, base_dir=out_dir)
                        if not os.path.isabs(img_path) or not os.path.exists(img_path):
                            if out_dir:
                                candidate = os.path.normpath(os.path.abspath(os.path.join(out_dir, src_url)))
                                if os.path.exists(candidate) and os.path.isfile(candidate):
                                    img_path = candidate
                            else:
                                base_filename = os.path.basename(src_url)
                                if os.path.exists(asset_mgr.cache_dir):
                                    for root, _, files in os.walk(asset_mgr.cache_dir):
                                        if base_filename in files:
                                            cand = os.path.join(root, base_filename)
                                            if os.path.isfile(cand):
                                                img_path = cand
                                                break
                        
                        img_path = os.path.normpath(os.path.abspath(img_path))
                        if os.path.exists(img_path) and os.path.isfile(img_path):
                            try:
                                p = doc.add_paragraph()
                                # Apply alignment from token
                                if tok.align == "center":
                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                elif tok.align == "right":
                                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                else:
                                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                                run = p.add_run()
                                
                                from PIL import Image as PILImage
                                with PILImage.open(img_path) as pil_img:
                                    w_px, h_px = pil_img.size

                                # Check for custom width (% or px)
                                if tok.width:
                                    if tok.width.endswith("%"):
                                        try:
                                            pct = float(tok.width.rstrip("%"))
                                            w_inches = min(6.0, max(0.5, 6.0 * (pct / 100.0)))
                                        except Exception:
                                            w_inches = min(5.8, max(1.5, w_px / 150.0))
                                    else:
                                        try:
                                            px_val = float(tok.width.rstrip("px"))
                                            w_inches = min(6.0, max(0.5, px_val / 96.0))
                                        except Exception:
                                            w_inches = min(5.8, max(1.5, w_px / 150.0))
                                else:
                                    w_inches = min(5.8, max(1.5, w_px / 150.0))
                                    
                                run.add_picture(img_path, width=Inches(w_inches))
                                p.paragraph_format.space_before = Pt(6)
                                p.paragraph_format.space_after = Pt(6)
                                i += 1
                                continue
                            except Exception as img_err:
                                print(f"[DEBUG] WordModule: Failed to insert picture {img_path}: {img_err}")

            # Heading check (Fast string pre-check)
            if line.startswith("#"):
                m = re.match(r"^(#{1,6})\s+(.*)", line)
                if m:
                    level = len(m.group(1))
                    heading = doc.add_heading('', level=min(level, 9))
                    add_formatted_runs(heading, m.group(2), size=HEADING_SIZES[level], default_bold=True, default_color=HEADING_COLORS[level])
                    heading.paragraph_format.space_before = Pt(10)
                    heading.paragraph_format.space_after = Pt(4)
                    i += 1
                    continue

            # Table check (Fast string pre-check)
            if "|" in line and not re.match(r"^[\|\s\-:]+$", line.strip()):
                start_i = i
                table_lines = []
                while i < len(lines) and "|" in lines[i]:
                    table_lines.append(lines[i].strip())
                    i += 1
                from src.core.converters import parse_table_rows
                rows = parse_table_rows(table_lines)
                if len(rows) >= 2:
                    max_cols = len(rows[0])
                    tbl = doc.add_table(rows=len(rows), cols=max_cols)
                    tbl.style = "Table Grid"
                    for r_idx, row_data in enumerate(rows):
                        for c_idx, cell_text in enumerate(row_data):
                            cell = tbl.cell(r_idx, c_idx)
                            cell.text = ""
                            clean_text = cell_text.replace("<br>", "\n").replace("<br/>", "\n")
                            lines_in_cell = clean_text.split("\n")
                            for l_idx, c_line in enumerate(lines_in_cell):
                                p = cell.paragraphs[0] if l_idx == 0 else cell.add_paragraph()
                                if r_idx == 0:
                                    add_formatted_runs(p, c_line, size=11, default_bold=True, default_color="FFFFFF")
                                else:
                                    add_formatted_runs(p, c_line, size=11)

                            if r_idx == 0:
                                tc = cell._tc
                                tcPr = tc.get_or_add_tcPr()
                                shd = OxmlElement("w:shd")
                                shd.set(qn("w:fill"), "4472C4")
                                shd.set(qn("w:color"), "auto")
                                shd.set(qn("w:val"), "clear")
                                tcPr.append(shd)
                    doc.add_paragraph()
                    continue
                else:
                    # Not a valid Markdown table! Reset index back to process line as standard text
                    i = start_i

            # Bullet List check
            sline = line.lstrip()
            if sline and sline[0] in ("-", "*", "+"):
                m = re.match(r"^(\s*)[-*+]\s+(.*)", line)
                if m:
                    indent_spaces = len(m.group(1))
                    level = (indent_spaces // 2) + 1
                    style_name = f"List Bullet {level}" if level > 1 else "List Bullet"
                    if style_name not in doc.styles:
                        style_name = "List Bullet"
                    add_paragraph_with_font(doc, m.group(2), style=style_name)
                    i += 1
                    continue

            # Numbered List check
            if sline and sline[0].isdigit():
                m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
                if m:
                    indent_spaces = len(m.group(1))
                    level = (indent_spaces // 2) + 1
                    style_name = f"List Number {level}" if level > 1 else "List Number"
                    if style_name not in doc.styles:
                        style_name = "List Number"
                    add_paragraph_with_font(doc, m.group(2), style=style_name)
                    i += 1
                    continue

            if sline.startswith(("---", "***", "___")) and re.match(r"^[-*_]{3,}$", line.strip()):
                p = doc.add_paragraph("─" * 50)
                p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                i += 1
                continue

            text = re.sub(r"`(.*?)`", r"\1", line)
            if text.strip():
                add_paragraph_with_font(doc, text)
            i += 1

        doc.save(out_path)
        return f"Word document created successfully -> {os.path.basename(out_path)}"

ModuleRegistry.register(WordModule())
